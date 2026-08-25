from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(r"E:\大学课程\竞赛\时察千机——工业多变量时序智能运维平台项目计划书.docx")

NAVY = "16324F"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRAY = "5B6573"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color="1F2933", bold=False, italic=False):
    run.font.name = "等线"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, after=6, before=0, line=1.25, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, size=11, color="1F2933", bold=False, italic=False,
             after=6, before=0, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, after, before, 1.25, align)
    r = p.add_run(text)
    set_font(r, size, color, bold, italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, after=7 if level == 1 else 5, before=17 if level == 1 else 10, line=1.15)
    r = p.add_run(text)
    if level == 1:
        set_font(r, 16, BLUE, True)
    elif level == 2:
        set_font(r, 13, NAVY, True)
    else:
        set_font(r, 11.5, NAVY, True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.2)
    r = p.add_run(text)
    set_font(r, 11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=4, line=1.2)
    r = p.add_run(text)
    set_font(r, 11)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.1)
        r = p.add_run(text)
        set_font(r, 10.5, NAVY, True)
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell.text = ""
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.15)
            r = p.add_run(text)
            set_font(r, 10.2)
    set_table_widths(table, widths)
    add_text(doc, "", size=2, after=2)
    return table


def set_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_font(run, 9, GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = paragraph.add_run(" 页")
    set_font(run2, 9, GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "等线"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, NAVY), ("Heading 3", 11.5, NAVY)):
        st = doc.styles[name]
        st.font.name = "等线"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("时察千机  |  项目计划书")
    set_font(hr, 9, GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_page_number(footer)


def build():
    doc = Document()
    configure_document(doc)

    add_text(doc, "项目计划书", size=14, color=BLUE, bold=True, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "时察千机——工业多变量时序智能运维平台", size=25, color=NAVY, bold=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "面向工业设备连续监测、异常研判与维修闭环的智能运维方案", size=13, color=GRAY, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(doc, ["项目属性", "内容"], [
        ("项目方向", "工业智能运维与多变量时序分析"),
        ("目标用户", "设备运维人员、生产管理人员、设备服务单位"),
        ("当前阶段", "公开数据机制验证与本地原型联调"),
        ("申报用途", "竞赛项目计划书"),
    ], [2100, 7260])
    add_text(doc, "说明：当前公开数据仅用于验证算法流程和工程闭环，相关实验结果不代表企业现场收益。", size=10, color=RED, after=24, before=4)
    add_heading(doc, "一、项目摘要")
    add_text(doc, "时察千机是一套面向工业设备多变量时序数据的智能运维平台。平台持续接收设备运行过程中产生的压力、流量、温度、振动、电流、阀门开度和控制状态等数据，在新数据批次到达后，自动完成数据质量检查、异常检测、趋势预测、多模型交叉验证、根因候选分析、风险分级、维修工单生成和消息通知。人员完成接单和现场处置后，系统依据同一数据源在处置后的新数据批次开展自动复检，判断异常是否消失，形成从监测、分析、告警、处置到验证的完整闭环。")
    add_text(doc, "项目以元景万悟作为自动化编排入口，以时察千机后端和关系数据库作为业务计算、数据存储与审计基础，以网页工作台展示分析证据和工单状态。算法计算、工单生成、通知、时限督办和维修复检等主链不依赖大语言模型，智能模型和知识资料主要用于辅助解释与问答，兼顾自动化能力、运行稳定性和过程可追溯性。")

    add_heading(doc, "二、行业痛点与需求分析")
    add_text(doc, "工业设备会持续产生多路相互关联的时间序列数据。传统运维方式通常依靠人工查看曲线、经验判断和电话或群聊派单，难以同时兼顾实时性、准确性和过程留痕。异常发现后，分析结论与维修任务之间经常存在断点，人员接单、处置时限和维修效果也缺少统一的跟踪机制。")
    for item in [
        "数据多而分散：设备同时产生多路测点，单独查看某一个测点容易漏掉变量之间的关系变化。",
        "异常研判依赖经验：告警出现后，人员需要在多条曲线、历史记录和设备资料之间反复比对。",
        "处置过程不连续：告警、派单、接单、处理和复检往往分散在不同工具中，难以形成完整记录。",
        "维修效果难验证：如果没有处置后的同源新数据，不能可靠判断设备是否恢复，容易出现误完成。",
        "结论缺少证据：模型结果、知识来源和人员操作没有统一保存，不利于复盘、交接和责任追踪。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、产品定位与目标用户")
    add_heading(doc, "三点一 产品定位", 2)
    add_text(doc, "本项目不是单次文件分析工具，也不是只展示告警的监控看板，而是把工业数据分析、运维决策、工单协同和维修后验证连接起来的连续运维平台。平台适用于具有连续采集数据、需要及时识别异常并组织处置的设备和生产环节。")
    add_heading(doc, "三点二 目标用户", 2)
    add_table(doc, ["用户", "核心需求", "平台价值"], [
        ("设备运维人员", "快速理解异常、接收任务、记录处置并确认复检", "提供证据、工单和后续验证结果"),
        ("生产管理人员", "掌握风险、超时任务和班次状态", "提供风险分级、时限督办和班次简报"),
        ("设备制造与服务单位", "积累设备故障案例和维修经验", "沉淀同源数据、分析证据和知识资料"),
        ("数据与算法人员", "验证模型、阈值和设备适配效果", "提供模型对比、指标记录和可审计结果"),
    ], [2100, 3300, 3960])
    add_heading(doc, "三点三 项目愿景", 2)
    add_text(doc, "形成可适配不同工业设备的数据驱动运维平台，使设备风险能够被更早发现，处置任务能够被更快组织，维修效果能够被数据验证，历史故障能够转化为可复用的运维知识。")

    add_heading(doc, "四、产品功能与应用场景")
    add_table(doc, ["功能模块", "主要能力", "典型使用场景"], [
        ("数据监测", "自动发现监测目录中的新数据批次并执行质量检查", "设备连续运行和定期批次采集"),
        ("异常分析", "识别异常点和异常事件，分析趋势并对比多种检测方法", "阀门、泵组和生产装置状态监测"),
        ("决策支持", "形成异常证据、根因候选、风险等级和处置建议", "运维人员研判告警和安排维修"),
        ("工单协同", "自动生成工单，支持接单、处理中、待验证和已完成状态", "班组任务分派与过程跟踪"),
        ("自动复检", "使用处置后的同源新数据验证异常是否消失", "维修完成后的效果确认"),
        ("辅助问答", "根据结构化证据和知识资料解释原因与处置依据", "交接班、复盘和运维知识查询"),
    ], [1900, 4100, 3360])

    add_heading(doc, "五、技术方案与系统架构")
    add_text(doc, "平台采用自动化编排入口、业务服务层、数据存储层和展示交互层协同工作的架构。各层职责清晰，核心工业流程由可复现的算法和业务规则执行，智能模型不作为异常判断、工单生成和复检完成的唯一依据。")
    add_table(doc, ["层次", "组成", "职责"], [
        ("自动化编排层", "元景万悟已发布工作流及定时触发器", "负责时间触发、工具编排、运行记录和结果输出"),
        ("业务服务层", "时察千机后端工业分析工具", "负责数据质量、异常检测、趋势预测、根因候选、工单和通知"),
        ("数据与审计层", "关系数据库及分析记录", "保存数据源、数据批次、分析任务、模型证据、工单和操作记录"),
        ("展示交互层", "网页运维工作台", "展示分析证据、工单状态、通知记录和闭环进度"),
        ("辅助解释层", "智能模型与工业知识资料", "解释结构化结果和提供知识问答，不替代核心业务判断"),
    ], [2100, 3600, 3660])
    add_heading(doc, "五点一 多变量时序分析", 2)
    add_text(doc, "多变量是指系统同时分析同一设备或工艺中的多路关联信号，而不是只判断一个指标是否越过阈值。例如阀门开度、流量和压力应具有一定的协同变化关系，单个数值尚未越限时，变量之间的关系异常也可能提示执行机构卡滞、泄漏或控制失效风险。")
    add_heading(doc, "五点二 多模型交叉验证", 2)
    add_text(doc, "平台保留主模型结果，并使用具有互补特点的检测方法进行独立对比，输出异常点数量、异常事件、模型一致性、适用条件和选择原因。交叉验证结果作为可信度证据展示，不简单用严格多数结果压制主模型告警，避免降低异常事件召回能力。")

    add_heading(doc, "六、自动化运维工作流设计")
    add_text(doc, "平台通过五条独立工作流形成分工明确的自动化闭环。数据源配置只用于首次接入或修改配置，不参与周期运行。其余四条工作流分别承担巡检、督办、复检和交接班汇总，互不混入不属于自身职责的业务。")
    add_table(doc, ["工作流", "触发方式", "核心动作", "输出"], [
        ("数据源接入配置", "首次配置或变更时手动触发", "登记监测目录、设备类型和数据约定", "可用数据源配置"),
        ("无人值守工业巡检", "新数据到达后的周期触发", "发现数据、分析、读取证据并分级通知", "分析任务、风险结果和告警"),
        ("工单时限督办", "按周期检查未接单和超时工单", "按风险等级提醒、催办和升级", "督办记录和升级通知"),
        ("维修后自动复检", "按周期检查待验证工单", "查找处置后的同源成功批次并复检", "完成或退回处理的结果"),
        ("工业班次简报", "按班次周期汇总运行情况", "统计任务、异常、工单、督办、复检和通知", "班次简报"),
    ], [2200, 2300, 3360, 1500])
    add_heading(doc, "六点一 运维状态闭环", 2)
    add_number(doc, "新数据进入监测目录，系统自动登记数据批次并执行质量检查。")
    add_number(doc, "无人值守巡检完成异常检测、趋势研判和多模型结果对比。")
    add_number(doc, "系统生成风险分级、根因候选和维修工单，并向对应人员发送通知。")
    add_number(doc, "人员确认接单后，将工单更新为处理中，填写根因确认和处置记录。")
    add_number(doc, "处置完成后将工单设置为待验证，系统等待同一数据源的新成功批次。")
    add_number(doc, "维修后复检通过则自动完成，未通过则退回处理中并再次通知。")

    add_heading(doc, "七、项目创新点")
    for item in [
        "从一次性分析转向连续闭环：把监测、分析、告警、工单、处置和复检统一为可追踪过程。",
        "从单点阈值转向多变量关联研判：综合多路时间序列变化、持续状态和变量关系识别风险。",
        "从单一模型结论转向多模型证据：同时展示不同方法的结果和一致性，便于人员复核。",
        "从大模型主导转向确定性主流程加智能辅助：核心业务不依赖模型接口稳定性，模型用于解释和问答。",
        "从告警结束转向维修效果验证：没有处置后的新数据时保持等待，避免误判设备恢复。",
        "从结果展示转向过程审计：保存数据批次、分析证据、知识来源、工单操作和通知状态，支持复盘。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "八、市场分析与竞争差异")
    add_text(doc, "本项目的竞争对象主要包括传统人工巡检方式、只提供单次文件分析的工具、只展示告警的监控看板以及依赖人工派单和人工复核的运维系统。这些方案在单点监测或数据展示方面各有价值，但通常难以同时完成多变量分析、工单协同、时限督办和维修后验证。")
    add_table(doc, ["比较维度", "传统方式或单一工具", "时察千机"], [
        ("异常识别", "人工经验或单一阈值", "多变量时序分析与多模型证据"),
        ("告警处置", "电话、群聊或人工派单", "自动生成工单并按角色通知"),
        ("超时管理", "依赖管理人员记忆", "独立工作流自动提醒和升级"),
        ("维修验证", "人工查看或缺少复检", "同源新批次自动复检，结果回写工单"),
        ("结果解释", "经验说明或孤立报告", "结构化证据结合知识资料辅助解释"),
        ("过程追踪", "记录分散", "数据、模型、工单和通知统一审计"),
    ], [2200, 3380, 3780])

    add_heading(doc, "九、当前验证情况与应用边界")
    add_text(doc, "当前项目已完成本地环境下的自动巡检、异常分析、工单生成、消息通知、人员处置和维修后自动复检联调，并使用公开工业数据集完成机制验证。公开数据可以用于展示系统如何发现异常、形成证据和完成闭环，但不能据此宣称已经取得企业现场收益，也不能替代真实设备的工程验收。")
    add_text(doc, "现阶段仍需在真实设备接入前重新确认字段、单位、采样周期、健康基线、告警阈值、故障标签和安全处置边界。对于可能影响生产安全的控制建议，平台只提供辅助信息，必须由具备权限的人员确认后执行。")
    add_heading(doc, "九点一 已完成验证", 2)
    for item in [
        "自动发现监测目录中的新数据批次，并跳过已处理的重复内容。",
        "完成异常检测、趋势预测、根因候选和多模型结果对比。",
        "自动生成维修工单并支持确认接单、处理中、待验证和已完成状态。",
        "支持同一数据源处置后的新批次自动复检，通过或未通过均能回写工单。",
        "支持独立的时限督办、复检和班次简报工作流。",
        "支持网页工作台查看分析证据、工单状态和操作记录。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十、实施计划与发展路径")
    add_table(doc, ["阶段", "重点工作", "阶段成果"], [
        ("第一阶段", "完成公开数据集上的算法、工作流和闭环验证", "可演示的本地原型"),
        ("第二阶段", "完善设备配置、知识资料、模型证据和审计展示", "可复用的设备适配方案"),
        ("第三阶段", "接入试点设备数据，重新标定阈值和健康基线", "面向试点的运行版本"),
        ("第四阶段", "完善组织权限、通知渠道、确认机制和安全边界", "可管理的运维应用"),
        ("第五阶段", "沉淀设备类型模板和故障案例，形成标准化部署方案", "可推广的产品化方案"),
    ], [1800, 4800, 2760])

    add_heading(doc, "十一、风险控制与安全边界")
    for item in [
        "数据风险：真实设备接入前确认字段、单位、采样周期和数据质量，避免直接套用公开数据参数。",
        "模型风险：模型结果必须与结构化证据和规则校验结合，不能让自然语言输出直接改变工单状态。",
        "处置风险：涉及生产控制、停机或参数调整的建议必须由授权人员确认，并保留回退条件。",
        "平台风险：周期工作流分别执行，核心流程不依赖大模型接口，减少限流和外部服务波动影响。",
        "隐私风险：数据库连接信息、平台密钥和通知地址只保存在本地环境配置中，不写入计划书和代码仓库。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "十二、竞赛演示方案")
    add_number(doc, "启动本地数据库、自动化编排平台、业务后端、网页工作台和四条周期触发器。")
    add_number(doc, "向演示监测目录投放一份新的公开工业数据，展示无人值守巡检运行记录。")
    add_number(doc, "在网页工作台查看异常证据、模型结果对比、趋势和根因候选。")
    add_number(doc, "查看自动生成的维修工单和分级通知，模拟人员确认接单并填写处置记录。")
    add_number(doc, "将工单设置为待验证，投放同一数据源的后续数据，展示自动复检结果。")
    add_number(doc, "在自动化编排平台查看巡检、督办、复检和班次简报的独立运行记录。")
    add_text(doc, "演示时应明确说明：公开数据用于机制验证，演示结果不代表企业现场收益；真实部署需要重新完成设备适配和工程验收。", size=10, color=RED, before=6)

    add_heading(doc, "十三、结语")
    add_text(doc, "时察千机的核心价值不是让用户在聊天框中上传文件后获得一次性分析，而是让工业设备数据进入持续运行的自动化运维流程。平台将多变量时序分析转化为可解释的风险证据，将风险证据转化为可执行的维修任务，再通过同源新数据验证处置效果。通过自动化编排、确定性业务服务、网页运维工作台和智能辅助解释的协同，项目具备进一步接入真实设备、完善安全机制和形成行业化方案的基础。")

    doc.core_properties.title = "时察千机——工业多变量时序智能运维平台项目计划书"
    doc.core_properties.subject = "工业智能运维竞赛项目计划书"
    doc.core_properties.author = "时察千机项目组"
    doc.core_properties.comments = ""
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
