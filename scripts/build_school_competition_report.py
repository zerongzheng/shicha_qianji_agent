from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"E:\大学课程\竞赛")
PROJECT = ROOT / "shicha_qianji_agent"
ASSETS = PROJECT / ".runtime-temp" / "school-report-assets"
DOCX_OUT = ROOT / "理工科类赛道+时察千机——工业多变量时序智能运维平台+项目报告.docx"

NAVY = "17324D"
TEAL = "176B73"
LIGHT_TEAL = "E8F2F1"
LIGHT_BLUE = "EAF0F6"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "64727A"
RED = "9B1C1C"
BLACK = "000000"
TABLE_WIDTH = 9020


OVERVIEW = (
    "工业设备持续产生压力、流量、温度、振动、电流和控制状态等多变量时序数据。传统方式多依赖人工查看曲线和固定阈值，"
    "异常发现、原因研判、维修派单与效果验证相互割裂。现有异常检测研究通常聚焦单次算法输出，通用智能体又容易把数值判断交给"
    "不可复现的自然语言推理，难以满足工业场景对稳定性、证据链和责任闭环的要求。\n"
    "本作品研发“时察千机”工业多变量时序智能运维平台，以元景万悟承担自动化编排，以确定性时序算法和业务规则承担数据质量检查、"
    "异常检测、趋势预测、根因候选、工单生成、通知、督办和维修后复检。系统自动发现新数据批次，不要求用户先上传文件再发起问答；"
    "智能模型与知识库只用于辅助解释，不改变已落库的分析和工单事实。平台形成“新数据到达—自动分析—风险告警—人员接单—现场处置"
    "—同源新数据复检”的闭环，面向阀门、泵组和连续生产装置等场景。当前使用公开SKAB数据集进行机制验证，后续接入真实设备时仅需"
    "替换数据适配、健康基线和设备知识，并重新校准工程阈值。"
)

DESIGN = (
    "系统采用分层架构。元景万悟负责五条已发布工作流：数据源接入配置、无人值守工业巡检、工单时限督办、维修后自动复检和工业班次"
    "简报；外部触发器只提供时间信号。FastAPI后端暴露十九个稳定工具，执行数据去重、不可变快照、异步任务、算法分析和业务动作；"
    "PostgreSQL保存数据源、分析批次、模型证据、智能体决策、工单、通知与审计记录；Vue3工作台展示风险、执行链、决策账本、趋势预测和"
    "模型证据。辅助解释与知识库向量检索使用阿里云百炼模型服务，并记录模型名称、调用状态、耗时和令牌统计；接口异常不影响确定性主链。\n"
    "数据进入后，系统依据时间轴、期望采样间隔和缺失模式自动完成时间对齐、保守填补与质量闸门；再面向多变量任务构建局部残差、一阶变化率、"
    "局部波动及传感器关系特征，生成滑动窗口。归一化采用健康基线或当前文件前段训练窗口拟合RobustScaler，后续分析点只执行变换，避免"
    "未来信息进入当前判断。原始快照和预处理动作均保留，质量门失败时阻断自动分析。\n"
    "模型路由不读取当前文件异常标签，而是综合分析目标、设备配置、数据规模、传感器数量和健康基线可用性，对六类检测器按冻结能力顺序进行"
    "适用性排序，选择当前任务场景下最适配的主模型，并同步确定阈值与事件策略；首选模型不满足最低条件或运行失败时自动回退。主模型从时域"
    "偏离、局部频域变化和测点关系变化构造风险分数，其他模型比较事件数、异常点和一致性。严格多数共识会降低事件召回，因此交叉验证作为"
    "可信度证据而不抑制主告警。\n"
    "趋势模块使用最近值、指数平滑、局部线性、滞后岭回归和时频增强岭回归，通过无泄漏滚动回测选择模型并输出预测区间。根因模块比较事件"
    "前后测点变化、相关性与领先时滞，将结果匹配为待验证候选，并给出证据缺口和现场检查步骤。人员处置后将工单设为待验证；复检工作流只"
    "使用同一数据源在处置后的新成功批次，原异常消失则完成，仍存在则退回处理中，没有新数据则保持等待。核心链路不依赖大模型，模型或"
    "知识库异常时不会阻断巡检和处置。"
)

TESTING = (
    "测试环境为Windows 11、Python 3.13.12和PostgreSQL，本地以Docker运行元景万悟基础服务。软件测试固定项目内临时目录执行，"
    "后端共166项测试通过，前端30项测试通过，覆盖自适应预处理、场景化模型选择、六类检测器、交叉验证、趋势预测、优化建议、决策账本、"
    "异步任务、数据仓储、身份权限、工单状态、时限督办、同源复检、知识检索、模型审计和万悟接口；Ruff静态检查通过。只读接入验收显示"
    "后端与数据库就绪，完整OpenAPI包含十九个工具，四条周期工作流"
    "的发布标识和密钥配置均有效，验收过程未投递样本或发送通知。\n"
    "算法实验使用公开SKAB数据集，按完整CSV文件划分为一个健康基线、十七个验证文件和十七个独立测试文件，避免拆分连续序列造成时序"
    "泄漏。阈值、最短事件长度和合并间隔只在验证集选择，冻结后在独立测试集形成六个检测器共102条记录。时频关系多路径检测器事件级F1"
    "为0.6196、事件召回为0.9412、点级F1为0.3433、平均误报事件为每文件1.47个，单文件平均耗时0.2237秒；稳健基线平均误报事件"
    "为1.41个。四模型三票严格多数共识事件召回降至0.8235，故未部署为主告警门槛。\n"
    "趋势预测在十七个测试文件的136条传感器序列上全部成功，滚动回测自动选模相对持续模型的标准化误差改善4.75%，百分之九十五区间"
    "覆盖率为95.12%。固定随机种子的四类受控退化场景均在越界前预警，平均提前20.75个采样点；这些模拟阈值不是设备工程限值。工作流"
    "联调已验证新批次自动分析、工单生成、人员处置、待验证和同源新数据复检。当前结果只证明公开数据上的算法与工程机制，不代表企业现场"
    "收益。"
)

INNOVATION = (
    "作品的原创性体现在数据自适应、模型路由、算法证据与智能体审计的协同设计。第一，系统不是用固定脚本直接套用单一模型，而是自动记录"
    "时间对齐、缺失填补、特征构建、窗口生成和防泄漏归一化证据；再根据任务目标、设备配置、数据规模、传感器数量和健康基线选择当前场景"
    "最适配的主模型，并保留候选排序、选择原因和回退条件。第二，主模型联合时域偏离、频域变化和测点关系变化，六类检测器并列交叉验证。"
    "项目保留“严格多数共识未优于主模型”的负结果，避免为了形式上的多模型而牺牲事件召回。第三，系统把每次关键判断整理为智能体决策账本，"
    "记录触发条件、结构化证据、冻结规则、执行动作、人工闸门与回退条件；趋势、根因和历史知识形成的优化建议同时给出建议范围、观察窗口、"
    "验证指标和回退条件，且不直接控制设备。第四，采用确定性主链与智能辅助分离架构。数值计算、工单、通知、督办和复检不依赖大模型；"
    "检索增强与语言模型仅解释结构化证据，并记录知识来源、模型名称、调用状态、耗时和令牌统计。五条工作流按职责拆分，巡检主链不混入"
    "督办和复检，周期任务通过发布时预授权运行。\n"
    "作品具备本地可部署、可测试和可扩展特征。一键脚本统一检查数据库、容器平台、后端、网页工作台和四条周期触发器，并使用独立进程标识"
    "防止重复启动。数据适配、设备配置和知识资料与算法主线解耦，可面向阀门、泵组等设备扩展。其现实价值在于把分散的监测、研判、派单和"
    "复检连接为可追溯流程，帮助运维人员聚焦证据和现场验证；但涉及停机、控制参数和安全联锁的建议必须由授权人员确认，平台不直接下发"
    "控制指令。"
)

ETHICS = (
    "作品坚持数据最小化、可追溯和人工负责原则。当前仅使用公开SKAB数据进行机制验证，并在报告中明确数据来源和用途；不采集人脸、语音、"
    "地图等敏感信息，不调用境外接口。企业接入前须确认数据权属、脱敏要求、保存期限和访问权限，并重新建立健康基线及独立测试。数据库连接、"
    "平台密钥和通知地址仅存放在被版本控制忽略的本地环境文件中，不写入源码、报告或前端。\n"
    "异常标签、趋势风险和根因输出均为辅助证据。根因置信度表示候选的现场核查优先级，不解释为故障发生概率；预测结果不能替代设备安全限值"
    "和停机规程。涉及通知、工单变更和数据源配置的操作保留身份、岗位和审计记录，关键处置由人员确认。检索增强只接收测点、候选原因和证据"
    "缺口等最小上下文，大模型不可直接读取整份原始数据或修改业务事实。作品由团队原创实现，使用的开源平台、框架、模型方法和公开数据均在"
    "参考文献及说明中标注，不修改元景万悟平台源码。"
)

SUMMARY = (
    "时察千机围绕工业多变量时序数据构建了可运行、可测试、可审计的智能运维闭环。系统以元景万悟作为自动化编排入口，以确定性算法完成"
    "自适应数据处理、场景化模型路由、异常检测、多模型证据、趋势预测和根因候选，以智能体决策账本记录证据、规则、动作与回退条件，并由"
    "PostgreSQL和网页工作台承接工单、通知、督办、人员处置与维修后复检。"
    "公开SKAB独立测试验证了主模型在事件发现和误报控制之间的平衡，全量软件测试与只读平台验收验证了系统工程完整性。\n"
    "作品不把大语言模型包装成工业判断核心，而是将其限制在辅助解释和知识问答范围，从而保留算法复现、故障降级和人工确认能力。下一阶段"
    "将接入真实设备数据与工况日志，重新校准阈值、预测区间和设备专属规则，完善多设备并发、模型回滚和组织权限，并在授权范围内开展独立"
    "时间段试点验证。"
)


def section_count(text: str) -> int:
    return len("".join(text.split()))


def assert_limits() -> None:
    limits = {
        "作品概述": (OVERVIEW, 600),
        "作品设计与实现": (DESIGN, 1000),
        "作品测试与分析": (TESTING, 1000),
        "创新性与实用性说明": (INNOVATION, 800),
        "伦理与合规性说明": (ETHICS, 500),
    }
    for name, (text, limit) in limits.items():
        count = section_count(text)
        if count > limit:
            raise ValueError(f"{name} exceeds {limit}: {count}")
        print(f"{name}: {count}/{limit}")


def chinese_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc")
    return ImageFont.truetype(str(path), size=size)


def rounded_box(draw, xy, text, fill, outline, font, text_fill=BLACK, radius=20):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    left, top, right, bottom = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    x = (left + right - (bbox[2] - bbox[0])) / 2
    y = (top + bottom - (bbox[3] - bbox[1])) / 2
    normalized_text_fill = text_fill if str(text_fill).startswith("#") else f"#{text_fill}"
    draw.multiline_text(
        (x, y),
        text,
        font=font,
        fill=normalized_text_fill,
        spacing=8,
        align="center",
    )


def arrow(draw, start, end, fill=TEAL, width=5):
    draw.line((start, end), fill=f"#{fill}", width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 15
    base_x, base_y = x2 - ux * size, y2 - uy * size
    draw.polygon(
        [(x2, y2), (base_x + px * 7, base_y + py * 7), (base_x - px * 7, base_y - py * 7)],
        fill=f"#{fill}",
    )


def build_architecture_image() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "architecture.png"
    image = Image.new("RGB", (1800, 860), "white")
    draw = ImageDraw.Draw(image)
    title_font = chinese_font(46, True)
    body_font = chinese_font(27, True)
    small_font = chinese_font(23)
    draw.text((70, 42), "时察千机系统架构与责任边界", font=title_font, fill=f"#{NAVY}")
    draw.text((72, 105), "核心业务由确定性算法和规则执行，智能模型只做辅助解释", font=small_font, fill=f"#{MID_GRAY}")

    boxes = [
        ((70, 210, 380, 365), "工业数据源\n目录 / HTTP", LIGHT_GRAY, MID_GRAY),
        ((510, 175, 870, 400), "元景万悟\n五条独立工作流\n定时触发与编排", LIGHT_TEAL, TEAL),
        ((1000, 175, 1380, 400), "时察千机后端\n19 个工业工具\n分析、路由与业务规则", LIGHT_BLUE, NAVY),
        ((1490, 210, 1730, 365), "PostgreSQL\n任务 / 工单\n审计", LIGHT_GRAY, MID_GRAY),
        ((1000, 545, 1380, 740), "Vue3 运维工作台\n证据 / 趋势 / 工单", LIGHT_TEAL, TEAL),
        ((510, 545, 870, 740), "知识库与语言模型\n辅助解释 / 问答\n故障时可降级", "FFF4DC", "A66A12"),
    ]
    for xy, text, fill, outline in boxes:
        rounded_box(draw, xy, text, f"#{fill}", f"#{outline}", body_font)
    arrow(draw, (380, 288), (510, 288))
    arrow(draw, (870, 288), (1000, 288))
    arrow(draw, (1380, 288), (1490, 288))
    arrow(draw, (1190, 400), (1190, 545))
    arrow(draw, (1000, 642), (870, 642))
    draw.text((70, 800), "数据链路：新批次 → 自动分析 → 风险证据 → 工单 → 处置 → 同源新数据复检", font=body_font, fill=f"#{NAVY}")
    image.save(out, optimize=True)
    return out


def build_workflow_image() -> Path:
    out = ASSETS / "workflow.png"
    image = Image.new("RGB", (1800, 620), "white")
    draw = ImageDraw.Draw(image)
    title_font = chinese_font(42, True)
    body_font = chinese_font(24, True)
    small_font = chinese_font(21)
    draw.text((70, 42), "无人值守运维闭环", font=title_font, fill=f"#{NAVY}")
    labels = [
        "新数据\n到达",
        "巡检分析\n多模型证据",
        "工单与\n分级通知",
        "人员接单\n现场处置",
        "待验证\n等待新数据",
        "自动复检\n完成或退回",
    ]
    lefts = [55, 345, 635, 925, 1215, 1505]
    for index, (label, left) in enumerate(zip(labels, lefts)):
        fill = LIGHT_TEAL if index in (1, 5) else LIGHT_GRAY
        outline = TEAL if index in (1, 5) else MID_GRAY
        rounded_box(draw, (left, 190, left + 230, 365), label, f"#{fill}", f"#{outline}", body_font)
        if index < len(labels) - 1:
            arrow(draw, (left + 230, 278), (lefts[index + 1], 278))
    draw.text((83, 430), "没有处置后的同源新成功批次时，工单保持“待验证”，系统不误判设备恢复。", font=small_font, fill=f"#{RED}")
    draw.text((83, 485), "周期工作流：无人值守巡检 60 秒；工单督办 300 秒；维修复检 300 秒；班次简报 8 小时。", font=small_font, fill=f"#{MID_GRAY}")
    image.save(out, optimize=True)
    return out


def build_model_chart() -> Path:
    out = ASSETS / "model_comparison.png"
    models = ["时频关系", "稳健MAD", "时序工况", "PCA重构", "自编码器", "孤立森林"]
    event_f1 = [0.6196, 0.5647, 0.4386, 0.4329, 0.4240, 0.2982]
    recall = [0.9412, 0.8824, 0.9412, 0.8824, 0.9412, 0.8824]
    false_events = [1.47, 1.41, 3.47, 3.06, 4.00, 9.65]
    colors = ["#176B73", "#4B9B99", "#8AA7A6", "#A9B7C5", "#74C7BB", "#C8D0D7"]
    image = Image.new("RGB", (2200, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = chinese_font(44, True)
    axis_font = chinese_font(22)
    label_font = chinese_font(23, True)
    note_font = chinese_font(20)
    draw.text((590, 35), "六类异常检测器对比（17个独立测试文件）", font=title_font, fill=f"#{NAVY}")

    left_origin = (110, 710)
    left_width, chart_height = 930, 520
    draw.line((left_origin[0], 190, left_origin[0], left_origin[1]), fill="#9AA7AD", width=3)
    draw.line((left_origin[0], left_origin[1], left_origin[0] + left_width, left_origin[1]), fill="#9AA7AD", width=3)
    for tick in range(6):
        y = left_origin[1] - round(chart_height * tick / 5)
        draw.line((left_origin[0], y, left_origin[0] + left_width, y), fill="#E5E9EC", width=2)
        draw.text((45, y - 14), f"{tick / 5:.1f}", font=axis_font, fill=f"#{MID_GRAY}")
    group_width = left_width / len(models)
    for index, model in enumerate(models):
        center = left_origin[0] + group_width * (index + 0.5)
        f1_h = chart_height * event_f1[index]
        recall_h = chart_height * recall[index]
        draw.rectangle((center - 43, left_origin[1] - f1_h, center - 5, left_origin[1]), fill=colors[index])
        draw.rectangle((center + 5, left_origin[1] - recall_h, center + 43, left_origin[1]), fill="#F0B45A")
        bbox = draw.textbbox((0, 0), model, font=axis_font)
        draw.text((center - (bbox[2] - bbox[0]) / 2, 735), model, font=axis_font, fill=f"#{BLACK}")
    draw.text((340, 125), "事件识别能力", font=label_font, fill=f"#{NAVY}")
    draw.rectangle((250, 805, 285, 833), fill=f"#{TEAL}")
    draw.text((300, 803), "事件级F1", font=note_font, fill=f"#{BLACK}")
    draw.rectangle((470, 805, 505, 833), fill="#F0B45A")
    draw.text((520, 803), "事件召回", font=note_font, fill=f"#{BLACK}")

    right_origin = (1210, 710)
    right_width = 850
    draw.line((right_origin[0], 190, right_origin[0], right_origin[1]), fill="#9AA7AD", width=3)
    draw.line((right_origin[0], right_origin[1], right_origin[0] + right_width, right_origin[1]), fill="#9AA7AD", width=3)
    for tick in range(6):
        y = right_origin[1] - round(chart_height * tick / 5)
        draw.line((right_origin[0], y, right_origin[0] + right_width, y), fill="#E5E9EC", width=2)
        draw.text((1150, y - 14), f"{tick * 2}", font=axis_font, fill=f"#{MID_GRAY}")
    group_width = right_width / len(models)
    for index, model in enumerate(models):
        center = right_origin[0] + group_width * (index + 0.5)
        bar_h = chart_height * false_events[index] / 10
        draw.rectangle((center - 28, right_origin[1] - bar_h, center + 28, right_origin[1]), fill=colors[index])
        bbox = draw.textbbox((0, 0), model, font=axis_font)
        draw.text((center - (bbox[2] - bbox[0]) / 2, 735), model, font=axis_font, fill=f"#{BLACK}")
    draw.text((1410, 125), "平均误报事件（个/文件）", font=label_font, fill=f"#{NAVY}")
    draw.text((760, 858), "参数仅在验证集选择，测试集不参与调参。", font=note_font, fill=f"#{MID_GRAY}")
    image.save(out, optimize=True)
    return out


def set_run_font(run, name="宋体", size=12, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, before, after in (
        ("Heading 1", 18, 18, 12),
        ("Heading 2", 15, 12, 6),
        ("Heading 3", 13, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("时察千机  |  理工科类赛道项目报告")
    set_run_font(run, size=9, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    tail = footer.add_run(" 页")
    set_run_font(tail, size=9, color=MID_GRAY)


def add_body(doc, text, first_indent=True, after=0):
    for raw in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(24 if first_indent else 0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(raw)
        set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_picture(doc, path, width_cm, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        p.text = ""
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(0)
            p.text = ""
            run = p.add_run(str(text))
            set_run_font(run, size=12)
    set_table_geometry(table, widths)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)
    run.add_break(WD_BREAK.PAGE)


def build_report():
    assert_limits()
    architecture = build_architecture_image()
    workflow = build_workflow_image()
    chart = build_model_chart()
    screenshot = PROJECT / "outputs" / "frontend_analysis_chain_final.png"
    wanwu_screenshot = ASSETS / "wanwu_agent_config.png"
    logo = ROOT / "时察千机项目标志.png"

    doc = Document()
    configure_document(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(16)
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Cm(2.8))
    for text, size, after in (
        ("行业智能体创新大赛", 22, 8),
        ("理工科类", 28, 4),
        ("项目报告", 28, 34),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        set_run_font(run, name="黑体", size=size, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("时察千机——工业多变量时序智能运维平台")
    set_run_font(run, name="黑体", size=20, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(70)
    run = p.add_run("提交日期：2026年8月")
    set_run_font(run, size=14)

    page_break(doc)
    add_heading(doc, "填写说明", 1)
    instructions = [
        "所有参赛项目必须为一个基本完整的智能体应用设计方案。作品报告书旨在能够清晰准确地阐述（或图示）该参赛队的参赛项目（或方案）。",
        "作品报告采用A4纸撰写。除标题外，所有内容必需为宋体、小四号字、1.5倍行距。",
        "作品报告中各项目说明文字部分仅供参考，作品报告书撰写完毕后，请删除所有说明文字。(本页不删除)",
        "作品报告模板里已经列的内容仅供参考，作者可以在此基础上增加内容或对文档结构进行微调。",
        "为保证网评的公平、公正，作品报告中应避免出现作者所在学校、院系和指导教师等泄露身份的信息。一经发现，取消作品参赛资格。",
    ]
    for index, item in enumerate(instructions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{index}. {item}")
        set_run_font(run)

    page_break(doc)
    add_heading(doc, "目  录", 1)
    toc_rows = [
        ("一、作品概述", "4"),
        ("二、作品设计与实现", "5"),
        ("三、作品测试与分析", "8"),
        ("四、创新性与实用性说明", "10"),
        ("五、伦理与合规性说明", "12"),
        ("六、总结", "13"),
        ("七、参考文献", "14"),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [7700, 1320])
    for title, page in toc_rows:
        cells = table.add_row().cells
        for cell in cells:
            cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
        p = cells[0].paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(title)
        set_run_font(run, size=13)
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(page)
        set_run_font(run, size=13)

    page_break(doc)
    add_heading(doc, "一、作品概述", 1)
    add_body(doc, OVERVIEW)
    add_picture(doc, workflow, 16.0, "图1  无人值守工业运维闭环")

    page_break(doc)
    add_heading(doc, "二、作品设计与实现", 1)
    add_body(doc, DESIGN)
    add_picture(doc, architecture, 16.0, "图2  系统分层架构与责任边界")
    add_table_caption(doc, "表1  五条元景万悟工作流的职责划分")
    add_table(
        doc,
        ["工作流", "触发方式", "主要职责"],
        [
            ("数据源接入配置", "首次配置或变更", "登记目录、设备类型和数据约定"),
            ("无人值守工业巡检", "每60秒", "发现新批次、分析、读取证据并通知"),
            ("工单时限督办", "每300秒", "未接单提醒和超时升级"),
            ("维修后自动复检", "每300秒", "检查待验证工单和处置后同源新批次"),
            ("工业班次简报", "每8小时", "汇总任务、异常、工单、复检和通知"),
        ],
        [2500, 1800, 4720],
    )
    if wanwu_screenshot.exists():
        add_picture(doc, wanwu_screenshot, 14.2, "图3  万悟平台智能体配置与工作流工具编排界面（自动化主入口）")

    page_break(doc)
    add_heading(doc, "三、作品测试与分析", 1)
    add_body(doc, TESTING)
    add_table_caption(doc, "表2  六类检测器独立测试结果（17个测试文件）")
    add_table(
        doc,
        ["检测器", "事件级F1", "事件召回", "点级F1", "误报事件/文件", "耗时/秒"],
        [
            ("时频关系多路径", "0.6196", "0.9412", "0.3433", "1.47", "0.2237"),
            ("稳健MAD", "0.5647", "0.8824", "0.1234", "1.41", "0.1973"),
            ("时序工况混合", "0.4386", "0.9412", "0.2383", "3.47", "0.8460"),
            ("PCA多变量重构", "0.4329", "0.8824", "0.2258", "3.06", "0.1768"),
            ("滑动窗口自编码器", "0.4240", "0.9412", "0.3198", "4.00", "0.1815"),
            ("孤立森林", "0.2982", "0.8824", "0.2007", "9.65", "0.6656"),
        ],
        [2350, 1350, 1350, 1250, 1550, 1170],
    )
    add_picture(doc, chart, 15.8, "图4  独立测试事件识别与误报对比")
    if screenshot.exists():
        add_picture(doc, screenshot, 16.0, "图5  自适应处理、场景化模型选择与可追溯执行链（16/16）")

    add_heading(doc, "四、创新性与实用性说明", 1)
    add_body(doc, INNOVATION)
    page_break(doc)
    add_table_caption(doc, "表3  作品创新点与传统方案差异")
    add_table(
        doc,
        ["创新维度", "传统方案", "本作品"],
        [
            ("数据治理", "固定清洗脚本，处理过程难追溯", "自适应对齐、填补、特征、窗口与防泄漏归一化"),
            ("模型路由", "固定模型或依赖人工选择", "按目标、设备和数据条件选择最适配主模型并可回退"),
            ("异常证据", "单阈值或单模型结论", "主模型加六模型对照和关系证据"),
            ("智能体审计", "只保留最终文本或告警", "决策账本记录证据、规则、动作、人工闸门和回退条件"),
            ("业务流程", "告警后依赖人工转述", "自动通知、督办、处置和同源复检"),
        ],
        [1750, 3250, 4020],
    )

    page_break(doc)
    add_heading(doc, "五、伦理与合规性说明", 1)
    add_body(doc, ETHICS)
    add_table_caption(doc, "表4  主要伦理与合规风险控制")
    add_table(
        doc,
        ["风险", "控制措施"],
        [
            ("数据权属与隐私", "公开数据标明来源；企业接入前确认授权、脱敏和保存期限"),
            ("模型误判", "并列展示证据与边界，根因候选必须由人员现场确认"),
            ("操作副作用", "通知、配置和工单变更需要权限或部署时预授权并留痕"),
            ("生产安全", "不直接下发控制指令，停机和参数调整由授权人员决定"),
            ("密钥安全", "密钥仅存本地环境配置，不进入源码、报告和前端"),
        ],
        [2450, 6570],
    )

    page_break(doc)
    add_heading(doc, "六、总结", 1)
    add_body(doc, SUMMARY)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("公开数据验证 · 确定性主链 · 人工确认 · 全过程审计")
    set_run_font(run, name="黑体", size=14, bold=True, color=TEAL)

    page_break(doc)
    add_heading(doc, "七、参考文献", 1)
    refs = [
        "[1] SKOLTECH. Skoltech Anomaly Benchmark (SKAB)[EB/OL]. https://github.com/waico/SKAB, 2026-08-18.",
        "[2] LIU F T, TING K M, ZHOU Z H. Isolation Forest[C]//Proceedings of the 2008 IEEE International Conference on Data Mining. Piscataway: IEEE, 2008: 413-422.",
        "[3] SAKURADA M, YAIRI T. Anomaly Detection Using Autoencoders with Nonlinear Dimensionality Reduction[C]//Proceedings of the MLSDA 2014 Workshop. New York: ACM, 2014: 4-11.",
        "[4] PEDREGOSA F, VAROQUAUX G, GRAMFORT A, et al. Scikit-learn: Machine Learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "[5] LAVIN A, AHMAD S. Evaluating Real-Time Anomaly Detection Algorithms: The Numenta Anomaly Benchmark[C]//2015 IEEE 14th International Conference on Machine Learning and Applications. Piscataway: IEEE, 2015: 38-44.",
        "[6] UNICOMAI. 元景万悟：一站式行业大模型开发平台[EB/OL]. https://gitee.com/unicomai/wanwu, 2026-08-18.",
        "[7] POSTGRESQL GLOBAL DEVELOPMENT GROUP. PostgreSQL Documentation[EB/OL]. https://www.postgresql.org/docs/, 2026-08-18.",
        "[8] FASTAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2026-08-18.",
        "[9] DOCKER INC. Docker Documentation[EB/OL]. https://docs.docker.com/, 2026-08-23.",
        "[10] VUE.JS TEAM. Vue.js Guide[EB/OL]. https://vuejs.org/guide/, 2026-08-23.",
        "[11] 阿里云. 百炼大模型服务平台文档[EB/OL]. https://help.aliyun.com/zh/model-studio/, 2026-08-23.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.hanging_indent = Pt(24)
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        set_run_font(run, size=12)

    doc.core_properties.title = "时察千机——工业多变量时序智能运维平台项目报告"
    doc.core_properties.subject = "理工科类赛道匿名评审项目报告"
    doc.core_properties.author = "参赛团队"
    doc.core_properties.keywords = "工业时序,异常检测,智能运维,行业智能体"
    doc.core_properties.comments = ""
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build_report()
